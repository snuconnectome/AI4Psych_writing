#!/usr/bin/env python3
"""
Convert intro_complete_revised.txt to intro_complete_revised.docx
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Read the complete revised introduction
with open('/Users/jeongyeon/AI4Psych_writing/group/jeongyeonshin/intro_complete_revised.txt', 'r') as f:
    content = f.read()

# Create a new Word document
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Add title
title = doc.add_heading('Introduction', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.LEFT

# Split content into paragraphs
paragraphs = content.strip().split('\n\n')

# Section headings to identify
section_headings = [
    'Early life stress and neurocognitive functions',
    'Reward processing',
    'Emotion processing',
    'Working memory',
    'Impulsivity',
    'Research gap and objectives'
]

# Add each paragraph
for para_text in paragraphs:
    if para_text.strip():
        # Check if it's a section heading
        if para_text.strip() in section_headings:
            # Add as heading
            heading = doc.add_heading(para_text.strip(), level=2)
        else:
            # Add as paragraph
            p = doc.add_paragraph(para_text.strip())
            p_format = p.paragraph_format
            p_format.line_spacing = 2.0  # Double spacing
            p_format.space_after = Pt(0)
            p_format.first_line_indent = Inches(0.5)  # Indent first line
            p_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Justify text

# Save the document
output_path = '/Users/jeongyeon/AI4Psych_writing/group/jeongyeonshin/intro_complete_revised.docx'
doc.save(output_path)

print(f"✅ Complete Introduction Word document created: {output_path}")
print(f"📄 Total paragraphs: {len(paragraphs)}")
print(f"📊 Word count: {len(content.split())} words")
